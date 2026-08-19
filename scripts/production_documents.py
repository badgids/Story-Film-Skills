#!/usr/bin/env python3
# Copyright 2026 Alan Guice (Badgids)
# SPDX-License-Identifier: Apache-2.0
from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from copy import copy

from media_runtime import MediaRuntimeError, portable_rel, project_path, project_root

DOC_RX = re.compile(r'^DOC-\d{3}$')


def read_json(path: Path):
    return json.loads(path.read_text(encoding='utf-8'))


def load_manifest(path: Path) -> dict:
    if not path.exists():
        raise MediaRuntimeError(f'missing document manifest: {path}')
    obj = read_json(path)
    if obj.get('schema_version') != 1 or not isinstance(obj.get('documents'), list):
        raise MediaRuntimeError('document manifest requires schema_version=1 and documents array')
    return obj


def validate_manifest(root: Path, obj: dict) -> list[str]:
    errors=[]; seen=set()
    for i, rec in enumerate(obj.get('documents', []), 1):
        did=rec.get('doc_id',''); prefix=did or f'document {i}'
        if not DOC_RX.fullmatch(did or ''): errors.append(f'{prefix}: doc_id must be DOC-###')
        elif did in seen: errors.append(f'{did}: duplicate doc_id')
        seen.add(did)
        fmt=rec.get('format')
        if fmt not in {'xlsx','docx','pdf'}: errors.append(f'{prefix}: format must be xlsx, docx, or pdf')
        out=rec.get('path','')
        if not portable_rel(out): errors.append(f'{prefix}: path must be project-relative')
        else:
            try: project_path(root,out)
            except Exception as exc: errors.append(f'{prefix}: {exc}')
        md=rec.get('markdown_path') or (str(Path(out).with_suffix('.md')).replace('\\','/') if isinstance(out,str) and out else '')
        if not isinstance(md,str) or not portable_rel(md) or not md.lower().endswith('.md'):
            errors.append(f'{prefix}: markdown_path must be a project-relative .md companion')
        src=rec.get('data_path') or rec.get('content_path')
        if not isinstance(src,str) or not portable_rel(src): errors.append(f'{prefix}: data_path or content_path must be project-relative')
        sources=rec.get('source_ids',[])
        if not isinstance(sources,list): errors.append(f'{prefix}: source_ids must be an array')
    return errors


def set_cell_style(cell, *, bold=False, fill=None, color=None, size=None):
    from openpyxl.styles import Font, PatternFill
    cell.font = Font(name='Arial', bold=bold, color=color or '000000', size=size or 10)
    if fill:
        cell.fill = PatternFill('solid', fgColor=fill)


def render_xlsx(title: str, spec: dict, output: Path) -> None:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment
    wb=Workbook(); default=wb.active; wb.remove(default)
    sheets=spec.get('sheets',[])
    if not sheets: raise MediaRuntimeError('XLSX data spec requires sheets')
    for sidx,spec_sheet in enumerate(sheets):
        name=str(spec_sheet.get('name') or f'Sheet{sidx+1}')[:31]
        ws=wb.create_sheet(name)
        columns=spec_sheet.get('columns',[]); rows=spec_sheet.get('rows',[])
        if title and sidx==0:
            ws.cell(1,1,title); set_cell_style(ws.cell(1,1),bold=True,size=14)
            start=3
        else: start=1
        if columns:
            for c,head in enumerate(columns,1):
                cell=ws.cell(start,c,str(head)); set_cell_style(cell,bold=True,fill='D9EAF7'); cell.alignment=Alignment(horizontal='center')
            start += 1
        for r,row in enumerate(rows,start):
            values = row if isinstance(row,list) else [row.get(col,'') for col in columns]
            for c,val in enumerate(values,1):
                cell=ws.cell(r,c,val)
                set_cell_style(cell)
                if isinstance(val,str) and val.startswith('='):
                    font = copy(cell.font); font.color = '000000'; cell.font = font
        for c in range(1,max(1,len(columns)) + 1):
            maxlen=0
            for cell in ws.iter_cols(min_col=c,max_col=c):
                for x in cell:
                    if x.value is not None: maxlen=max(maxlen,len(str(x.value)))
            ws.column_dimensions[ws.cell(1,c).column_letter].width=min(max(maxlen+2,10),45)
        ws.freeze_panes = f'A{start}' if columns else None
    output.parent.mkdir(parents=True,exist_ok=True); wb.save(output)


def parse_markdown(text: str):
    blocks=[]
    for raw in text.splitlines():
        line=raw.rstrip()
        if not line: blocks.append(('blank','')); continue
        if line.startswith('### '): blocks.append(('h3',line[4:])); continue
        if line.startswith('## '): blocks.append(('h2',line[3:])); continue
        if line.startswith('# '): blocks.append(('h1',line[2:])); continue
        if line.startswith('- '): blocks.append(('bullet',line[2:])); continue
        m=re.match(r'^\d+\.\s+(.*)$',line)
        if m: blocks.append(('number',m.group(1))); continue
        blocks.append(('p',line))
    return blocks


def render_docx(title: str, text: str, output: Path) -> None:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.section import WD_SECTION
    doc=Document(); sec=doc.sections[0]; sec.top_margin=Inches(.65); sec.bottom_margin=Inches(.65); sec.left_margin=Inches(.7); sec.right_margin=Inches(.7)
    styles=doc.styles
    styles['Normal'].font.name='Arial'; styles['Normal'].font.size=Pt(10)
    if title: doc.add_heading(title,0)
    for kind,value in parse_markdown(text):
        if kind=='blank': continue
        if kind=='h1': doc.add_heading(value,1)
        elif kind=='h2': doc.add_heading(value,2)
        elif kind=='h3': doc.add_heading(value,3)
        elif kind=='bullet': doc.add_paragraph(value,style='List Bullet')
        elif kind=='number': doc.add_paragraph(value,style='List Number')
        else: doc.add_paragraph(value)
    footer=sec.footer.paragraphs[0]; footer.text='Generated from Story-Film Skills project sources'; footer.style=styles['Normal']
    output.parent.mkdir(parents=True,exist_ok=True); doc.save(output)


def render_pdf(title: str, text: str, output: Path) -> None:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_LEFT
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
    from reportlab.lib.units import inch
    from xml.sax.saxutils import escape
    styles=getSampleStyleSheet()
    styles['Normal'].fontName='Helvetica'; styles['Normal'].fontSize=9.5; styles['Normal'].leading=13
    doc=SimpleDocTemplate(str(output),pagesize=LETTER,rightMargin=.65*inch,leftMargin=.65*inch,topMargin=.65*inch,bottomMargin=.65*inch)
    story=[]
    if title: story.append(Paragraph(escape(title),styles['Title'])); story.append(Spacer(1,8))
    bullets=[]
    def flush_bullets():
        nonlocal bullets
        if bullets:
            story.append(ListFlowable([ListItem(Paragraph(escape(x),styles['Normal'])) for x in bullets],bulletType='bullet'))
            bullets=[]
    for kind,value in parse_markdown(text):
        if kind=='bullet': bullets.append(value); continue
        flush_bullets()
        if kind=='blank': story.append(Spacer(1,5))
        elif kind=='h1': story.append(Paragraph(escape(value),styles['Heading1']))
        elif kind=='h2': story.append(Paragraph(escape(value),styles['Heading2']))
        elif kind=='h3': story.append(Paragraph(escape(value),styles['Heading3']))
        elif kind=='number': story.append(Paragraph('• '+escape(value),styles['Normal']))
        else: story.append(Paragraph(escape(value),styles['Normal']))
    flush_bullets(); output.parent.mkdir(parents=True,exist_ok=True); doc.build(story)


def qc(path: Path, fmt: str) -> dict:
    if not path.exists() or path.stat().st_size == 0: return {'status':'fail','reason':'missing-or-empty'}
    if fmt=='xlsx':
        from openpyxl import load_workbook
        wb=load_workbook(path,read_only=True,data_only=False); return {'status':'pass','sheets':wb.sheetnames}
    if fmt=='docx':
        from zipfile import ZipFile
        with ZipFile(path) as z:
            names=set(z.namelist())
            if 'word/document.xml' not in names: return {'status':'fail','reason':'missing-document-xml'}
        return {'status':'pass'}
    if fmt=='pdf':
        from pypdf import PdfReader
        reader=PdfReader(str(path)); return {'status':'pass','pages':len(reader.pages)}
    return {'status':'fail','reason':'unsupported-format'}



def markdown_for_xlsx(title: str, spec: dict) -> str:
    lines=[f'# {title}' if title else '# Production Spreadsheet','']
    for sheet in spec.get('sheets',[]):
        name=str(sheet.get('name') or 'Sheet')
        columns=[str(x) for x in sheet.get('columns',[])]
        rows=sheet.get('rows',[])
        lines += [f'## {name}','']
        normalized=[]
        for row in rows:
            if isinstance(row,list): vals=row
            elif isinstance(row,dict): vals=[row.get(c,'') for c in columns]
            else: vals=[row]
            normalized.append(['' if v is None else str(v) for v in vals])
        width=max([len(columns),*(len(r) for r in normalized)],default=0)
        if width:
            headers=(columns+['']*width)[:width] if columns else [f'Column {i}' for i in range(1,width+1)]
            lines += ['| '+' | '.join(headers)+' |','| '+' | '.join(['---']*width)+' |']
            for row in normalized:
                row=(row+['']*width)[:width]; lines.append('| '+' | '.join(row)+' |')
        else:
            lines.append('_Empty sheet._')
        lines.append('')
    return '\n'.join(lines).rstrip()+'\n'

def markdown_companion_path(rec: dict) -> str:
    return rec.get('markdown_path') or str(Path(rec['path']).with_suffix('.md')).replace('\\','/')

def render_record(root: Path, rec: dict) -> dict:
    fmt=rec['format']; out=project_path(root,rec['path']); src_rel=rec.get('data_path') or rec.get('content_path'); src=project_path(root,src_rel,must_exist=True); title=str(rec.get('title',''))
    md_rel=markdown_companion_path(rec); md_out=project_path(root,md_rel)
    if fmt=='xlsx':
        spec=read_json(src); render_xlsx(title,spec,out); md_text=markdown_for_xlsx(title,spec)
    else:
        text=src.read_text(encoding='utf-8')
        if fmt=='docx': render_docx(title,text,out)
        elif fmt=='pdf': render_pdf(title,text,out)
        md_text=text if text.lstrip().startswith('#') else ((f'# {title}\n\n' if title else '')+text)
    md_out.parent.mkdir(parents=True,exist_ok=True); md_out.write_text(md_text.rstrip()+'\n',encoding='utf-8')
    result=qc(out,fmt); result.update({'doc_id':rec['doc_id'],'path':rec['path'],'format':fmt,'markdown_path':md_rel,'markdown_exists':md_out.is_file()}); return result


def main() -> int:
    ap=argparse.ArgumentParser(description='Render production XLSX, DOCX, and PDF deliverables from a project manifest.')
    ap.add_argument('project_dir'); ap.add_argument('--manifest',default='00_project/document_manifest.json'); ap.add_argument('--doc-id'); ap.add_argument('--validate-only',action='store_true')
    args=ap.parse_args(); root=project_root(args.project_dir); mp=project_path(root,args.manifest,must_exist=True); obj=load_manifest(mp); errors=validate_manifest(root,obj)
    if errors:
        for e in errors: print('ERROR',e)
        return 2
    rows=[r for r in obj['documents'] if not args.doc_id or r.get('doc_id')==args.doc_id]
    if args.doc_id and not rows: raise MediaRuntimeError(f'unknown doc_id {args.doc_id}')
    if args.validate_only:
        print(json.dumps({'status':'pass','documents':len(rows)},indent=2)); return 0
    results=[render_record(root,r) for r in rows]
    failed=[r for r in results if r['status']!='pass']
    print(json.dumps({'status':'fail' if failed else 'pass','results':results},indent=2)); return 2 if failed else 0


if __name__=='__main__':
    try: raise SystemExit(main())
    except MediaRuntimeError as exc:
        print(f'ERROR {exc}'); raise SystemExit(2)
