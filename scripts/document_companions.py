#!/usr/bin/env python3
# Copyright 2026 Alan Guice (Badgids)
# SPDX-License-Identifier: Apache-2.0
from __future__ import annotations
import argparse,json,shutil,subprocess,tempfile,zipfile,re
from pathlib import Path
from xml.etree import ElementTree as ET

REQUIRED={'.pdf','.docx','.dotx','.xlsx','.xlsm','.xltx','.pptx','.pptm','.odt','.ods','.odp'}
def companion(path:Path)->Path:return path.with_suffix('.md')
def xlsx_text(path:Path)->str:
 from openpyxl import load_workbook
 wb=load_workbook(path,read_only=True,data_only=False,keep_vba=path.suffix.lower()=='.xlsm')
 out=[f'# {path.stem}','']
 for ws in wb.worksheets:
  out += [f'## Sheet: {ws.title}','']
  rows=[[('' if c is None else str(c)) for c in row] for row in ws.iter_rows(values_only=True)]
  width=max((len(r) for r in rows),default=0)
  if width:
   rows=[r+['']*(width-len(r)) for r in rows]
   out += ['| '+' | '.join(rows[0])+' |','| '+' | '.join(['---']*width)+' |']
   out += ['| '+' | '.join(r)+' |' for r in rows[1:]]
  else: out += ['_Empty sheet._']
  out.append('')
 return '\n'.join(out).rstrip()+'\n'
def docx_text(path:Path)->str:
 from docx import Document
 d=Document(path); out=[f'# {path.stem}','']
 for p in d.paragraphs:
  if p.text.strip():out.append(p.text)
 for ti,t in enumerate(d.tables,1):
  out += ['',f'## Table {ti}','']
  rows=[[c.text.replace('\n',' ') for c in row.cells] for row in t.rows]
  if rows:
   w=max(len(r) for r in rows); rows=[r+['']*(w-len(r)) for r in rows]
   out += ['| '+' | '.join(rows[0])+' |','| '+' | '.join(['---']*w)+' |']+['| '+' | '.join(r)+' |' for r in rows[1:]]
 return '\n'.join(out).rstrip()+'\n'
def pdf_text(path:Path)->str:
 try:
  from pypdf import PdfReader
  r=PdfReader(str(path)); out=[f'# {path.stem}','']
  for i,p in enumerate(r.pages,1):out += [f'## Page {i}','',p.extract_text() or '_No extractable text on this page._','']
  return '\n'.join(out).rstrip()+'\n'
 except Exception as exc:return f'# {path.stem}\n\nPDF text extraction failed: {exc}\n'
def office_text(path:Path)->str:
 if shutil.which('pandoc'):
  p=subprocess.run(['pandoc','-t','gfm',str(path)],text=True,stdout=subprocess.PIPE,stderr=subprocess.PIPE)
  if p.returncode==0 and p.stdout.strip():return f'# {path.stem}\n\n'+p.stdout.strip()+'\n'
 ext=path.suffix.lower()
 try:
  with zipfile.ZipFile(path) as z:
   names=z.namelist()
   if ext in {'.pptx','.pptm'}:
    slides=sorted([n for n in names if re.fullmatch(r'ppt/slides/slide\d+\.xml',n)], key=lambda n:int(re.search(r'(\d+)',Path(n).name).group(1)))
    out=[f'# {path.stem}','']
    for i,n in enumerate(slides,1):
     root=ET.fromstring(z.read(n)); text=' '.join((x.text or '') for x in root.iter() if x.tag.endswith('}t')).strip()
     out += [f'## Slide {i}','',text or '_No extractable slide text._','']
    return '\n'.join(out).rstrip()+'\n'
   if ext in {'.odt','.ods','.odp'} and 'content.xml' in names:
    root=ET.fromstring(z.read('content.xml')); chunks=[]
    for x in root.iter():
     if x.text and x.text.strip():chunks.append(x.text.strip())
    return f'# {path.stem}\n\n'+'\n\n'.join(chunks)+'\n'
 except Exception:
  pass
 raise RuntimeError(f'cannot create a meaningful Markdown companion for {path.name} on this runtime; provide the authoritative Markdown source instead')
def generate(path:Path,out:Path|None=None)->Path:
 if not path.is_file():raise FileNotFoundError(path)
 ext=path.suffix.lower(); target=out or companion(path)
 if ext in {'.xlsx','.xlsm','.xltx'}:text=xlsx_text(path)
 elif ext in {'.docx','.dotx'}:text=docx_text(path)
 elif ext=='.pdf':text=pdf_text(path)
 else:text=office_text(path)
 target.write_text(text,encoding='utf-8'); return target
def audit(root:Path):
 missing=[]
 for p in root.rglob('*'):
  if p.is_file() and p.suffix.lower() in REQUIRED and not companion(p).is_file():missing.append(p.relative_to(root).as_posix())
 return missing
def main():
 ap=argparse.ArgumentParser(); sub=ap.add_subparsers(dest='cmd',required=True)
 p=sub.add_parser('generate');p.add_argument('file');p.add_argument('--out')
 p=sub.add_parser('audit');p.add_argument('project_dir');p.add_argument('--generate-missing',action='store_true')
 a=ap.parse_args()
 if a.cmd=='generate':print(generate(Path(a.file).expanduser().resolve(),Path(a.out).expanduser().resolve() if a.out else None));return 0
 r=Path(a.project_dir).expanduser().resolve();miss=audit(r)
 if a.generate_missing:
  for rel in list(miss): generate(r/rel)
  miss=audit(r)
 print(json.dumps({'status':'pass' if not miss else 'fail','missing_companions':miss},indent=2));return 0 if not miss else 2
if __name__=='__main__':raise SystemExit(main())
